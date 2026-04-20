"""DOCX adapter — per-paragraph and per-table-cell extraction.

Preserves document-level structure (sections, styles, tables) by mutating
the source document in place on write. Per-paragraph granularity is the
right coarseness for a translator — finer run-level splitting loses
context, coarser section-level loses formatting preservation.

Run-level formatting survives via the "first run wins" strategy: the
translated text replaces the first run's text and later runs are emptied.
Paragraphs with no runs (e.g. empty) are skipped.
"""

from __future__ import annotations

from collections.abc import Iterable
from pathlib import Path

from docx import Document

from ...domain import TranslatedUnit, Unit, UnitKind
from ...use_cases.ports import DocumentAdapter


class DocxAdapter(DocumentAdapter):
    extension = ".docx"

    def extract(self, source_path: Path) -> list[Unit]:
        doc = Document(str(source_path))
        units: list[Unit] = []
        idx = 0
        for p_i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            if not text.strip():
                continue
            units.append(
                Unit(
                    id=f"{idx:04d}",
                    kind=UnitKind.PARAGRAPH,
                    text=text,
                    meta={"kind": "paragraph", "index": p_i},
                )
            )
            idx += 1
        for t_i, table in enumerate(doc.tables):
            for r_i, row in enumerate(table.rows):
                for c_i, cell in enumerate(row.cells):
                    for cp_i, paragraph in enumerate(cell.paragraphs):
                        text = paragraph.text
                        if not text.strip():
                            continue
                        units.append(
                            Unit(
                                id=f"{idx:04d}",
                                kind=UnitKind.PARAGRAPH,
                                text=text,
                                meta={
                                    "kind": "table_cell",
                                    "table": t_i,
                                    "row": r_i,
                                    "col": c_i,
                                    "paragraph": cp_i,
                                },
                            )
                        )
                        idx += 1
        return units

    def write(
        self,
        source_path: Path,
        translated: Iterable[TranslatedUnit],
        target_lang: str,
        output_path: Path,
    ) -> None:
        doc = Document(str(source_path))
        for unit in translated:
            kind = unit.meta.get("kind")
            if kind == "paragraph":
                idx = int(unit.meta.get("index", -1))
                if 0 <= idx < len(doc.paragraphs):
                    _replace_paragraph_text(doc.paragraphs[idx], unit.target_text)
            elif kind == "table_cell":
                t_i = int(unit.meta.get("table", -1))
                r_i = int(unit.meta.get("row", -1))
                c_i = int(unit.meta.get("col", -1))
                cp_i = int(unit.meta.get("paragraph", -1))
                if not (0 <= t_i < len(doc.tables)):
                    continue
                table = doc.tables[t_i]
                if not (0 <= r_i < len(table.rows)):
                    continue
                row = table.rows[r_i]
                if not (0 <= c_i < len(row.cells)):
                    continue
                cell = row.cells[c_i]
                if not (0 <= cp_i < len(cell.paragraphs)):
                    continue
                _replace_paragraph_text(cell.paragraphs[cp_i], unit.target_text)
        doc.save(str(output_path))


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""
